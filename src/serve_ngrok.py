import argparse
import json
import logging
import os
import sys
from pathlib import Path
from typing import Optional

# Ensure the parent directory is in the path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn
from pyngrok import ngrok, conf

from src.rag_pipeline import RAGPipeline

try:
    from data_pipeline import build_chunks, _clean_text, _chunk_text, CHUNK_TOKEN_TARGET
    from retriever import DYNAMIC_CHUNKS_PATH
except ModuleNotFoundError:
    from src.data_pipeline import build_chunks, _clean_text, _chunk_text, CHUNK_TOKEN_TARGET
    from src.retriever import DYNAMIC_CHUNKS_PATH

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define the FastAPI app
app = FastAPI(
    title="NUST Bank RAG API",
    description="RAG Pipeline served via Ngrok from Google Colab",
    version="1.0"
)

# Allow CORS for ngrok/frontend communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize pipeline lazily to save memory during startup/docs building
pipeline = None

class QueryRequest(BaseModel):
    query: str

class QueryResponse(BaseModel):
    answer: str

# ── Ingestion models ─────────────────────────────────────────────────────────

class IngestRequest(BaseModel):
    """
    Accepts either:
      - FAQ-style:  { product, question, answer }
      - Freeform:   { text }              (product/source optional)
    """
    product: Optional[str] = "Dynamic"
    source: Optional[str] = "dynamic"
    question: Optional[str] = None   # FAQ-style
    answer: Optional[str] = None     # FAQ-style
    text: Optional[str] = None       # Freeform / policy doc

class IngestResponse(BaseModel):
    message: str
    chunks_added: int
    total_dynamic_chunks: int
    total_index_vectors: int

@app.on_event("startup")
async def startup_event():
    global pipeline
    logger.info("Initializing RAG Pipeline...")
    pipeline = RAGPipeline()
    logger.info("RAG Pipeline initialized successfully.")

@app.post("/chat", response_model=QueryResponse)
async def chat(request: QueryRequest):
    if not pipeline:
        raise HTTPException(status_code=503, detail="Pipeline not initialized")
    
    try:
        answer = pipeline.answer(request.query)
        return QueryResponse(answer=answer)
    except Exception as e:
        logger.error(f"Error generating answer: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health():
    return {"status": "ok", "pipeline_loaded": pipeline is not None}


# ── Ingestion routes ──────────────────────────────────────────────────────────

@app.post("/ingest", response_model=IngestResponse)
async def ingest(request: IngestRequest):
    """
    Add a new document/FAQ entry to the live knowledge base.

    Two accepted formats:
    1. FAQ-style (question + answer):
       { "product": "Loans", "question": "...", "answer": "..." }

    2. Freeform text (policy doc, article, etc.):
       { "product": "Policies", "source": "dynamic", "text": "..." }
    """
    if not pipeline:
        raise HTTPException(status_code=503, detail="Pipeline not initialized")

    # ── Build a raw pair/chunk list that matches data_pipeline's format ──
    if request.question and request.answer:
        # FAQ-style: let build_chunks format and split it
        pair = {
            "source": request.source or "dynamic",
            "product": request.product or "Dynamic",
            "question": _clean_text(request.question),
            "answer": _clean_text(request.answer),
        }
        chunks = build_chunks([pair])

    elif request.text:
        # Freeform: chunk the raw text directly
        cleaned = _clean_text(request.text)
        words = cleaned.split()
        if len(words) <= CHUNK_TOKEN_TARGET:
            raw_chunks = [cleaned]
        else:
            raw_chunks = _chunk_text(cleaned)
        chunks = [
            {
                "text": c,
                "source": request.source or "dynamic",
                "product": request.product or "Dynamic",
            }
            for c in raw_chunks
        ]

    else:
        raise HTTPException(
            status_code=422,
            detail="Provide either (question + answer) for FAQ-style, or (text) for freeform ingestion."
        )

    if not chunks:
        raise HTTPException(status_code=422, detail="No valid chunks could be built from the provided input.")

    try:
        added = pipeline.retriever.add_documents(chunks)
    except Exception as e:
        logger.error(f"Ingestion error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

    # Count total persisted dynamic chunks
    total_dynamic = 0
    if DYNAMIC_CHUNKS_PATH.exists():
        with open(DYNAMIC_CHUNKS_PATH, encoding="utf-8") as f:
            total_dynamic = len(json.load(f))

    return IngestResponse(
        message=f"Successfully ingested {added} chunk(s) from '{request.product}'.",
        chunks_added=added,
        total_dynamic_chunks=total_dynamic,
        total_index_vectors=pipeline.retriever.index.ntotal,
    )


@app.get("/ingest/status")
async def ingest_status():
    """Returns counts of base vs dynamic chunks in the live index."""
    if not pipeline:
        raise HTTPException(status_code=503, detail="Pipeline not initialized")

    total_dynamic = 0
    dynamic_docs: list[dict] = []
    if DYNAMIC_CHUNKS_PATH.exists():
        with open(DYNAMIC_CHUNKS_PATH, encoding="utf-8") as f:
            dynamic_docs = json.load(f)
        total_dynamic = len(dynamic_docs)

    # Summarise unique sources ingested
    sources = list({d.get("product", "unknown") for d in dynamic_docs})

    return {
        "total_index_vectors": pipeline.retriever.index.ntotal,
        "dynamic_chunks": total_dynamic,
        "dynamic_products": sources,
        "dynamic_chunks_file": str(DYNAMIC_CHUNKS_PATH),
    }


@app.delete("/ingest/reset")
async def ingest_reset():
    """
    Wipes dynamic_chunks.json and reloads the retriever from the base index only.
    The base faiss.index is NEVER touched.
    """
    if not pipeline:
        raise HTTPException(status_code=503, detail="Pipeline not initialized")

    try:
        import faiss, pickle
        from src.retriever import INDEX_PATH, CHUNKS_STORE_PATH

        # Wipe the dynamic chunks file
        if DYNAMIC_CHUNKS_PATH.exists():
            DYNAMIC_CHUNKS_PATH.unlink()

        # Reload the retriever from base index (no dynamic chunks to merge)
        pipeline.retriever.index = faiss.read_index(str(INDEX_PATH))
        with open(CHUNKS_STORE_PATH, "rb") as f:
            pipeline.retriever.chunks = pickle.load(f)

        logger.info("Dynamic chunks reset. Retriever reloaded from base index.")
        return {
            "message": "Dynamic chunks wiped. Retriever restored to base index.",
            "total_index_vectors": pipeline.retriever.index.ntotal,
        }
    except Exception as e:
        logger.error(f"Reset error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ingest/file", response_model=IngestResponse)
async def ingest_file(
    file: UploadFile = File(...),
    product: str = Form(default="Dynamic"),
    source: str = Form(default="dynamic"),
):
    """
    Upload a Word (.docx) or Excel (.xlsx) file and ingest its text
    as dynamic chunks into the live knowledge base.

    Form fields:
      - file:    The uploaded .docx or .xlsx file
      - product: Label for the knowledge category (e.g. 'Home Loan Policy')
      - source:  Optional source tag (default: 'dynamic')
    """
    if not pipeline:
        raise HTTPException(status_code=503, detail="Pipeline not initialized")

    filename = file.filename or ""
    ext = Path(filename).suffix.lower()

    if ext not in (".docx", ".xlsx"):
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '{ext}'. Only .docx and .xlsx are accepted.",
        )

    content = await file.read()
    import io

    raw_texts: list[str] = []

    # ── Parse Word document ─────────────────────────────────────────────
    if ext == ".docx":
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="python-docx is not installed on this server.",
            )
        doc = DocxDocument(io.BytesIO(content))
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) > 10:
                raw_texts.append(text)
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if len(row_text) > 10:
                    raw_texts.append(row_text)

    # ── Parse Excel workbook ───────────────────────────────────────────
    elif ext == ".xlsx":
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows(values_only=True):
                row_text = " ".join(
                    str(c).strip() for c in row
                    if c is not None and str(c).strip()
                )
                row_text = _clean_text(row_text)
                if len(row_text) > 10:
                    raw_texts.append(row_text)

    if not raw_texts:
        raise HTTPException(
            status_code=422,
            detail="No readable text could be extracted from the uploaded file.",
        )

    # ── Chunk all extracted text ──────────────────────────────────────────
    full_text = "\n".join(raw_texts)
    cleaned = _clean_text(full_text)
    words = cleaned.split()
    if len(words) <= CHUNK_TOKEN_TARGET:
        split_chunks = [cleaned]
    else:
        split_chunks = _chunk_text(cleaned)

    chunks = [
        {
            "text": c,
            "source": source,
            "product": product,
        }
        for c in split_chunks
    ]

    try:
        added = pipeline.retriever.add_documents(chunks)
    except Exception as e:
        logger.error(f"File ingestion error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

    total_dynamic = 0
    if DYNAMIC_CHUNKS_PATH.exists():
        with open(DYNAMIC_CHUNKS_PATH, encoding="utf-8") as f:
            total_dynamic = len(json.load(f))

    logger.info(f"File '{filename}' ingested: {added} chunks added under product='{product}'")
    return IngestResponse(
        message=f"File '{filename}' ingested successfully: {added} chunk(s) added under '{product}'.",
        chunks_added=added,
        total_dynamic_chunks=total_dynamic,
        total_index_vectors=pipeline.retriever.index.ntotal,
    )


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Serve RAG Pipeline via ngrok")
    parser.add_argument("--port", type=int, default=8000, help="Port to run FastAPI on")
    parser.add_argument("--ngrok-token", type=str, required=True, help="Your ngrok authtoken")
    parser.add_argument("--region", type=str, default="us", help="ngrok region (e.g. us, eu, ap)")
    
    args = parser.parse_args()
    
    # Configure ngrok
    logger.info("Setting up ngrok...")
    conf.get_default().auth_token = args.ngrok_token
    conf.get_default().region = args.region
    
    # Open ngrok tunnel
    public_url = ngrok.connect(args.port).public_url
    logger.info(f"===============================================================")
    logger.info(f"🚀 Public URL: {public_url}")
    logger.info(f"API Docs available at: {public_url}/docs")
    logger.info(f"===============================================================")
    
    # Run uvicorn server
    uvicorn.run(app, host="0.0.0.0", port=args.port)
